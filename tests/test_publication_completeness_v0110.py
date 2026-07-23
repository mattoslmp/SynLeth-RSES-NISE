from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
import yaml

from rses_onco.publication import save_figure_triplet

ROOT = Path(__file__).resolve().parents[1]


def test_registry_has_complete_v0110_asset_contract() -> None:
  config = yaml.safe_load(
    (ROOT / "config/article_assets.yaml").read_text(encoding="utf-8")
  )
  assert {row["id"] for row in config["main_figures"]} == {
    f"Figure_{index}" for index in range(1, 9)
  }
  assert {row["id"] for row in config["supplementary_figures"]} == {
    f"Figure_S{index}" for index in range(1, 73)
  }
  assert len(config["main_tables"]) == 4
  assert len(config["supplementary_tables"]) == 47


def test_extended_figure_orchestrator_registers_s39_to_s69() -> None:
  source = (
    ROOT / "scripts/make_extended_supporting_figures.py"
  ).read_text(encoding="utf-8")
  assert "range(39, 70)" in source
  assert "wgcna_correlation_fallback_all_cancers.tsv" in source
  assert "Figure_S68" in (
    ROOT / "config/article_assets.yaml"
  ).read_text(encoding="utf-8")
  assert "Figure_S69" in (
    ROOT / "config/article_assets.yaml"
  ).read_text(encoding="utf-8")


def test_pipeline_exposes_extended_tables_and_document_build() -> None:
  source = "\n".join(
    (ROOT / path).read_text(encoding="utf-8")
    for path in (
      "scripts/publication_pipeline_steps.sh",
      "scripts/make_all_article_figures.py",
    )
  )
  for token in (
    "materialize_extended_supplementary_tables.py",
    "build_asset_reproduction_registry.py",
    "make_extended_supporting_figures.py",
    "build_publication_documents.py",
    "validate_publication_documents.py",
    "extended-tables",
    "documents",
  ):
    assert token in source


def test_figure_level_title_is_removed_before_export(tmp_path: Path) -> None:
  fig, axis = plt.subplots()
  fig.suptitle("General title that must not be embedded")
  axis.plot([0, 1], [0, 1])
  output = tmp_path / "figure"
  audit = save_figure_triplet(fig, output, "Figure_Test", strict_layout=False)
  assert audit.status in {"pass", "warning"}
  svg = output.with_suffix(".svg").read_text(encoding="utf-8")
  assert "General title that must not be embedded" not in svg


def test_document_builder_forces_s68_and_s69_page_separation() -> None:
  builder = (
    ROOT / "scripts/build_publication_documents.py"
  ).read_text(encoding="utf-8")
  validator = (
    ROOT / "scripts/validate_publication_documents.py"
  ).read_text(encoding="utf-8")
  assert "number in {68, 69}" in builder
  assert 'supplement_pages["Figure_S68"] == supplement_pages["Figure_S69"]' in validator


def test_reproduction_registry_is_a_supplementary_table() -> None:
  config = yaml.safe_load(
    (ROOT / "config/article_assets.yaml").read_text(encoding="utf-8")
  )
  assert "Table_S44_asset_reproduction_registry.tsv" in config["supplementary_tables"]
  registry = (
    ROOT / "scripts/build_asset_reproduction_registry.py"
  ).read_text(encoding="utf-8")
  assert "reproduction_command" in registry
  assert "document_location" in registry


def test_document_builder_creates_editable_docx_from_registered_assets(
  tmp_path: Path,
) -> None:
  import subprocess
  import sys
  from docx import Document
  from PIL import Image

  article_root = tmp_path / "article_outputs"
  figure_dir = article_root / "figures"
  manifest_dir = article_root / "manifests"
  manifest_dir.mkdir(parents=True)
  records = []
  for figure_id, category in (
    ("Figure_1", "main"),
    ("Figure_S68", "supplementary"),
    ("Figure_S69", "supplementary"),
  ):
    base = figure_dir / category / figure_id
    base.parent.mkdir(parents=True, exist_ok=True)
    Image.new("RGB", (120, 80), "white").save(base.with_suffix(".png"))
    records.append({
      "figure_id": figure_id,
      "category": category,
      "base_path": str(base),
      "caption": f"Caption for {figure_id}",
      "title": figure_id,
      "source_data_path": str(tmp_path / f"{figure_id}.tsv"),
    })
  pd.DataFrame(records).to_csv(
    manifest_dir / "figure_manifest.tsv", sep="\t", index=False
  )
  pd.DataFrame([
    {"table_id": "Table_1", "category": "main", "rows": 1, "columns": 1, "status": "ok", "path": "table1.tsv"},
    {"table_id": "Table_S1", "category": "supplementary", "rows": 1, "columns": 1, "status": "ok", "path": "tableS1.tsv"},
  ]).to_csv(manifest_dir / "table_manifest.tsv", sep="\t", index=False)
  manuscript = tmp_path / "manuscript.md"
  supplement = tmp_path / "supplement.md"
  manuscript.write_text("# Manuscript\nScientific text.\n", encoding="utf-8")
  supplement.write_text("# Supplement\nScientific text.\n", encoding="utf-8")
  output = article_root / "documents"
  subprocess.run([
    sys.executable,
    "scripts/build_publication_documents.py",
    "--article-root", str(article_root),
    "--manuscript-source", str(manuscript),
    "--supplement-source", str(supplement),
    "--output-dir", str(output),
    "--no-render-pdf",
  ], cwd=ROOT, check=True)
  main_doc = Document(output / "RSES_Onco_manuscript.docx")
  supplement_doc = Document(output / "RSES_Onco_supplementary_material.docx")
  assert len(main_doc.inline_shapes) == 1
  assert len(supplement_doc.inline_shapes) == 2
  assert sum(p.text.startswith("Figure_S68.") for p in supplement_doc.paragraphs) == 1
  assert sum(p.text.startswith("Figure_S69.") for p in supplement_doc.paragraphs) == 1
