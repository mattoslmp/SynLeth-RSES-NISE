#!/usr/bin/env python3
"""Build and optionally render the RSES-Onco manuscript and supplement.

The documents use only repository text sources, registered figure assets and
registered tables. Figure captions remain editable Word text. Supplementary
Figures S68 and S69 are forced onto separate pages.
"""
from __future__ import annotations

import argparse
from dataclasses import asdict, dataclass
from pathlib import Path
import re
import shutil
import subprocess

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]


@dataclass(frozen=True)
class DocumentRecord:
  document_id: str
  docx_path: str
  pdf_path: str
  source_path: str
  figures: int
  tables: int
  page_break_policy: str


def resolve_path(value: str | Path) -> Path:
  path = Path(value)
  return path if path.is_absolute() else ROOT / path


def add_page_field(paragraph) -> None:
  paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
  run = paragraph.add_run("Page ")
  begin = OxmlElement("w:fldChar")
  begin.set(qn("w:fldCharType"), "begin")
  instr = OxmlElement("w:instrText")
  instr.set(qn("xml:space"), "preserve")
  instr.text = " PAGE "
  separate = OxmlElement("w:fldChar")
  separate.set(qn("w:fldCharType"), "separate")
  text = OxmlElement("w:t")
  text.text = "1"
  end = OxmlElement("w:fldChar")
  end.set(qn("w:fldCharType"), "end")
  run._r.extend([begin, instr, separate, text, end])


def configure(document: Document) -> None:
  styles = document.styles
  styles["Normal"].font.name = "Arial"
  styles["Normal"].font.size = Pt(10.5)
  for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
    styles[name].font.name = "Arial"
  for section in document.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    add_page_field(section.footer.paragraphs[0])


def append_markdown(document: Document, path: Path) -> None:
  if not path.exists() or path.stat().st_size == 0:
    return
  for raw in path.read_text(encoding="utf-8").splitlines():
    line = raw.rstrip()
    if not line:
      document.add_paragraph()
    elif line.startswith("### "):
      document.add_heading(line[4:], level=3)
    elif line.startswith("## "):
      document.add_heading(line[3:], level=2)
    elif line.startswith("# "):
      document.add_heading(line[2:], level=1)
    elif line.startswith("- "):
      document.add_paragraph(line[2:], style="List Bullet")
    elif line.startswith("**") and line.endswith("**"):
      paragraph = document.add_paragraph()
      paragraph.add_run(line.strip("*")).bold = True
    elif line.startswith("|"):
      document.add_paragraph(line)
    else:
      document.add_paragraph(line)


def add_figure(document: Document, record: dict[str, object], page_break: bool) -> None:
  if page_break:
    document.add_page_break()
  figure_id = str(record["figure_id"])
  base = resolve_path(str(record["base_path"]))
  image = base.with_suffix(".png")
  if not image.exists() or image.stat().st_size == 0:
    raise FileNotFoundError(f"Missing figure image for document: {image}")
  paragraph = document.add_paragraph()
  paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
  paragraph.add_run().add_picture(str(image), width=Inches(6.6))
  caption = document.add_paragraph()
  caption.style = document.styles["Caption"]
  caption.add_run(f"{figure_id}. ").bold = True
  caption.add_run(str(record.get("caption") or record.get("title") or ""))
  source = document.add_paragraph()
  source.alignment = WD_ALIGN_PARAGRAPH.LEFT
  source.add_run("Exact figure source data: ").italic = True
  source.add_run(str(record.get("source_data_path") or "not recorded"))


def add_table_inventory(document: Document, tables: pd.DataFrame, category: str) -> None:
  subset = tables.loc[tables["category"].astype(str).eq(category)].copy()
  document.add_heading(
    "Main tables" if category == "main" else "Supplementary data tables",
    level=1,
  )
  table = document.add_table(rows=1, cols=5)
  table.style = "Table Grid"
  headers = ["Table", "Rows", "Columns", "Status", "Machine-readable file"]
  for cell, header in zip(table.rows[0].cells, headers):
    cell.text = header
  for record in subset.to_dict("records"):
    cells = table.add_row().cells
    values = [
      record.get("table_id"), record.get("rows"), record.get("columns"),
      record.get("status"), record.get("path"),
    ]
    for cell, value in zip(cells, values):
      cell.text = str(value)


def render_docx(docx: Path, output_dir: Path, libreoffice: str) -> Path:
  executable = shutil.which(libreoffice) or libreoffice
  profile = output_dir / ".libreoffice_profile"
  profile.mkdir(parents=True, exist_ok=True)
  subprocess.run([
    executable,
    f"-env:UserInstallation=file://{profile.resolve()}",
    "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx),
  ], cwd=ROOT, check=True)
  pdf = output_dir / f"{docx.stem}.pdf"
  if not pdf.exists() or pdf.stat().st_size == 0:
    raise RuntimeError(f"LibreOffice did not produce a non-empty PDF: {pdf}")
  return pdf


def render_pages(pdf: Path, output_dir: Path, pdftoppm: str) -> None:
  executable = shutil.which(pdftoppm) or pdftoppm
  page_dir = output_dir / f"{pdf.stem}_pages"
  if page_dir.exists():
    shutil.rmtree(page_dir)
  page_dir.mkdir(parents=True)
  subprocess.run([
    executable, "-png", "-r", "150", str(pdf), str(page_dir / "page")
  ], cwd=ROOT, check=True)
  if not list(page_dir.glob("page-*.png")):
    raise RuntimeError(f"No page renders were produced for {pdf}")


def main() -> None:
  parser = argparse.ArgumentParser()
  parser.add_argument("--article-root", default="article_outputs")
  parser.add_argument("--manuscript-source", default="manuscript/RSES_Onco_manuscript_source.md")
  parser.add_argument("--supplement-source", default="manuscript/RSES_Onco_supplement_source.md")
  parser.add_argument("--output-dir", default="article_outputs/documents")
  parser.add_argument("--render-pdf", action=argparse.BooleanOptionalAction, default=True)
  parser.add_argument("--render-pages", action=argparse.BooleanOptionalAction, default=False)
  parser.add_argument("--libreoffice", default="libreoffice")
  parser.add_argument("--pdftoppm", default="pdftoppm")
  args = parser.parse_args()

  article_root = resolve_path(args.article_root)
  output_dir = resolve_path(args.output_dir)
  output_dir.mkdir(parents=True, exist_ok=True)
  figures = pd.read_csv(article_root / "manifests/figure_manifest.tsv", sep="\t", low_memory=False)
  tables = pd.read_csv(article_root / "manifests/table_manifest.tsv", sep="\t", low_memory=False)

  main_figures = figures.loc[figures["category"].astype(str).eq("main")].copy()
  main_figures["number"] = main_figures["figure_id"].str.extract(r"(\d+)").astype(int)
  main_figures = main_figures.sort_values("number")
  supplement_figures = figures.loc[figures["category"].astype(str).eq("supplementary")].copy()
  supplement_figures["number"] = supplement_figures["figure_id"].str.extract(r"(\d+)").astype(int)
  supplement_figures = supplement_figures.sort_values("number")

  article = Document()
  configure(article)
  article.add_heading("RSES-Onco manuscript", level=0)
  append_markdown(article, resolve_path(args.manuscript_source))
  article.add_heading("Main figures", level=1)
  for index, record in enumerate(main_figures.to_dict("records")):
    add_figure(article, record, page_break=index > 0)
  add_table_inventory(article, tables, "main")
  article_path = output_dir / "RSES_Onco_manuscript.docx"
  article.save(article_path)

  supplement = Document()
  configure(supplement)
  supplement.add_heading("RSES-Onco supplementary material", level=0)
  append_markdown(supplement, resolve_path(args.supplement_source))
  methods_dir = article_root / "manuscript_assets/supplementary_methods"
  for path in sorted(methods_dir.glob("*.md")):
    append_markdown(supplement, path)
  supplement.add_heading("Supplementary figures", level=1)
  for index, record in enumerate(supplement_figures.to_dict("records")):
    number = int(record["number"])
    add_figure(supplement, record, page_break=index > 0 or number in {68, 69})
  add_table_inventory(supplement, tables, "supplementary")
  supplement_path = output_dir / "RSES_Onco_supplementary_material.docx"
  supplement.save(supplement_path)

  article_pdf = output_dir / "RSES_Onco_manuscript.pdf"
  supplement_pdf = output_dir / "RSES_Onco_supplementary_material.pdf"
  if args.render_pdf:
    article_pdf = render_docx(article_path, output_dir, args.libreoffice)
    supplement_pdf = render_docx(supplement_path, output_dir, args.libreoffice)
  if args.render_pages:
    render_pages(article_pdf, output_dir, args.pdftoppm)
    render_pages(supplement_pdf, output_dir, args.pdftoppm)

  records = [
    DocumentRecord(
      "main_manuscript", str(article_path), str(article_pdf), str(resolve_path(args.manuscript_source)),
      len(main_figures), int(tables["category"].astype(str).eq("main").sum()),
      "Main figures begin on separate pages.",
    ),
    DocumentRecord(
      "supplementary_material", str(supplement_path), str(supplement_pdf), str(resolve_path(args.supplement_source)),
      len(supplement_figures), int(tables["category"].astype(str).eq("supplementary").sum()),
      "Every supplementary figure begins on a separate page; S68 and S69 are explicitly separated.",
    ),
  ]
  pd.DataFrame([asdict(record) for record in records]).to_csv(
    output_dir / "document_build_manifest.tsv", sep="\t", index=False
  )
  print(f"Wrote {article_path}")
  print(f"Wrote {supplement_path}")


if __name__ == "__main__":
  main()
